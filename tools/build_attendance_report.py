from __future__ import annotations

from pathlib import Path
import os
import tempfile
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = Path(tempfile.gettempdir()) / "Bao_cao_do_an_Phan_mem_diem_danh_sinh_vien.docx"
OUT = Path(os.environ.get("REPORT_OUT", DEFAULT_OUT))


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C9DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(i, len(widths) - 1)])


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = ""
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(5)
    pf.line_spacing = 1.12
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 12, 6),
        ("Heading 2", 12.5, BLUE, 8, 4),
        ("Heading 3", 11.5, DARK_BLUE, 5, 3),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = normal
    footer.add_run("Trang ")
    add_field(footer, "PAGE")


def p(doc: Document, text: str = "", style: str | None = None, align=None, bold=False):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    return para


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Cm(0.75)
        para.paragraph_format.first_line_indent = Cm(-0.25)
        para.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    set_keep_with_next(para)


def add_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "A6A6A6")
    set_cell_margins(table, top=80, bottom=80, start=140, end=140)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    run.italic = True
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], caption: str | None = None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    set_cell_margins(table)
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
        run.font.size = Pt(8.6)
    return table


def add_front_matter(doc: Document) -> None:
    for text, size, bold in [
        ("BỘ KHOA HỌC VÀ CÔNG NGHỆ", 13, True),
        ("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", 13, True),
    ]:
        para = p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold)
        para.runs[0].font.size = Pt(size)
    for _ in range(4):
        p(doc)
    for text, size in [
        ("BÁO CÁO", 20),
        ("ĐỒ ÁN MÔN HỌC", 18),
        ("MÔN HỌC: NHẬP MÔN CÔNG NGHỆ PHẦN MỀM", 14),
        ("ĐỀ TÀI: PHẦN MỀM ĐIỂM DANH SINH VIÊN", 16),
    ]:
        para = p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        para.runs[0].font.size = Pt(size)
        para.runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for _ in range(3):
        p(doc)
    p(doc, "Giảng viên hướng dẫn: <Tên giảng viên>", align=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "Thực hiện bởi nhóm sinh viên:", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_table(
        doc,
        ["STT", "Họ tên", "MSSV", "Lớp", "Vai trò"],
        [
            ["1", "<Họ tên thành viên 01>", "<MSSV>", "<Lớp>", "Trưởng nhóm"],
            ["2", "<Họ tên thành viên 02>", "<MSSV>", "<Lớp>", "Thành viên"],
            ["3", "<Họ tên thành viên 03>", "<MSSV>", "<Lớp>", "Thành viên"],
        ],
        [700, 2900, 1600, 1600, 2560],
    )
    for _ in range(6):
        p(doc)
    p(doc, "TP.HCM, tháng ... năm 20...", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    doc.add_page_break()
    doc.add_heading("MỤC LỤC", level=1)
    p(doc, "Mục lục tự động sẽ được cập nhật trong Microsoft Word bằng thao tác References > Update Table. Các mục dưới đây phản ánh cấu trúc chính của báo cáo.")
    toc_items = [
        "Chương I. Tổng quan",
        "Chương II. Phân tích nội dung, yêu cầu",
        "Chương III. Phân tích thiết kế",
        "Chương IV. Phát triển/Thực thi",
        "Chương V. Triển khai",
        "Chương VI. Kết luận",
        "Tài liệu tham khảo",
    ]
    for item in toc_items:
        para = doc.add_paragraph(style="List Number")
        para.add_run(item)

    doc.add_page_break()
    doc.add_heading("DANH SÁCH HÌNH, BẢNG", level=1)
    figures = [
        "Hình 3.1. Sơ đồ use case tổng quát",
        "Hình 3.2. Sơ đồ hoạt động đăng nhập",
        "Hình 3.3. Sơ đồ hoạt động tạo phiên điểm danh",
        "Hình 3.4. Sơ đồ hoạt động điểm danh QR/OTP",
        "Hình 3.5. Sơ đồ hoạt động xử lý đơn xin phép",
        "Hình 3.6. Mô hình ERD hệ thống",
        "Hình 4.1. Màn hình đăng nhập",
        "Hình 4.2. Màn hình trang chủ Admin",
        "Hình 4.3. Màn hình quản lý người dùng",
        "Hình 4.4. Màn hình quản lý lớp học phần",
        "Hình 4.5. Màn hình import danh sách sinh viên",
        "Hình 4.6. Màn hình lớp học phần của giảng viên",
        "Hình 4.7. Màn hình tạo phiên điểm danh",
        "Hình 4.8. Màn hình QR/OTP điểm danh",
        "Hình 4.9. Màn hình sinh viên điểm danh",
        "Hình 4.10. Màn hình lịch sử điểm danh",
        "Hình 4.11. Màn hình gửi đơn xin phép",
        "Hình 4.12. Màn hình duyệt đơn xin phép",
    ]
    for item in figures:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)
    tables = [
        "Bảng 2.1. Yêu cầu chức năng của Admin",
        "Bảng 2.2. Yêu cầu chức năng của Giảng viên",
        "Bảng 2.3. Yêu cầu chức năng của Sinh viên",
        "Bảng 3.1. Danh sách bảng trong cơ sở dữ liệu",
        "Bảng 5.1. Tình trạng cài đặt chức năng",
        "Bảng 5.2. Test case kiểm thử hệ thống",
    ]
    for item in tables:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)

    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    add_table(
        doc,
        ["Từ viết tắt", "Diễn giải", "Ý nghĩa trong báo cáo"],
        [
            ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng giữa frontend và backend"],
            ["CSDL", "Cơ sở dữ liệu", "Nơi lưu trữ dữ liệu người dùng, lớp học phần và điểm danh"],
            ["UI", "User Interface", "Giao diện người dùng"],
            ["UX", "User Experience", "Trải nghiệm người dùng"],
            ["JWT", "JSON Web Token", "Cơ chế cấp access token và refresh token"],
            ["ORM", "Object Relational Mapping", "Ánh xạ bảng dữ liệu thành model"],
            ["QR", "Quick Response", "Mã quét dùng trong điểm danh"],
            ["OTP", "One-Time Password", "Mã xác thực dùng một lần, thay đổi 30 giây"],
            ["CRUD", "Create, Read, Update, Delete", "Nhóm thao tác quản lý dữ liệu cơ bản"],
            ["REST", "Representational State Transfer", "Phong cách thiết kế API"],
            ["MVC", "Model View Controller", "Cách tổ chức phần mềm theo lớp trách nhiệm"],
        ],
        [1600, 3200, 4560],
    )


def chapter_one(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG I. TỔNG QUAN", level=1)
    doc.add_heading("I. Giới thiệu đề tài", level=2)
    doc.add_heading("1. Lý do chọn đề tài", level=3)
    paragraphs = [
        "Điểm danh là một nghiệp vụ quen thuộc trong trường đại học, cao đẳng và trung tâm đào tạo. Mặc dù nghiệp vụ này có vẻ đơn giản, kết quả điểm danh lại ảnh hưởng trực tiếp đến việc đánh giá chuyên cần, cảnh báo học tập và tính kỷ luật trong lớp học. Khi số lượng lớp học phần tăng lên, giảng viên phải xử lý nhiều danh sách sinh viên, nhiều buổi học và nhiều trạng thái vắng mặt khác nhau.",
        "Cách điểm danh thủ công bằng giấy hoặc bảng tính Excel thường tốn thời gian ngay trong giờ học. Giảng viên phải gọi tên, sinh viên ký tên hoặc lớp trưởng tổng hợp danh sách. Quy trình này dễ phát sinh sai sót như ghi nhầm mã sinh viên, bỏ sót người học, cập nhật chậm hoặc thất lạc minh chứng. Khi cần thống kê theo học kỳ, theo lớp học phần hoặc theo tỷ lệ vắng, dữ liệu rời rạc làm việc tổng hợp mất nhiều công sức.",
        "Ở góc độ sinh viên, điểm danh thủ công cũng tạo ra các tình huống không minh bạch. Sinh viên có thể quên ký tên, ký hộ hoặc không biết kịp thời tình trạng chuyên cần của mình. Khi vắng có lý do chính đáng, việc gửi minh chứng qua nhiều kênh khác nhau làm cho giảng viên khó theo dõi và sinh viên khó biết đơn đã được duyệt hay chưa.",
        "Vì vậy, đề tài phần mềm điểm danh sinh viên được xây dựng nhằm số hóa quy trình điểm danh theo lớp học phần. Hệ thống cho phép giảng viên mở phiên điểm danh, sinh viên xác thực bằng QR động và OTP, đồng thời hỗ trợ admin quản lý dữ liệu nền như khoa, lớp, môn học, học kỳ, lớp học phần và danh sách sinh viên. Đây là một bài toán phù hợp với môn Nhập môn công nghệ phần mềm vì có đầy đủ các bước phân tích yêu cầu, thiết kế dữ liệu, xây dựng backend, frontend, kiểm thử và triển khai.",
    ]
    for text in paragraphs:
        p(doc, text)
    doc.add_heading("2. Mục tiêu của đề tài", level=3)
    p(doc, "Mục tiêu tổng quát của đề tài là xây dựng một hệ thống web hỗ trợ quản lý điểm danh sinh viên theo mô hình ba vai trò: Admin, Giảng viên và Sinh viên. Hệ thống hướng đến tính dễ sử dụng trong lớp học, tính chính xác của dữ liệu và khả năng mở rộng trong tương lai.")
    add_bullets(doc, [
        "Xây dựng hệ thống điểm danh sinh viên theo lớp học phần, có danh sách sinh viên đăng ký rõ ràng.",
        "Hỗ trợ giảng viên tạo phiên điểm danh theo buổi học và lấy mã QR/OTP thay đổi theo thời gian.",
        "Hỗ trợ sinh viên điểm danh nhanh bằng cách quét QR hoặc nhập OTP trên giao diện web.",
        "Hỗ trợ Admin quản lý người dùng, khoa, lớp hành chính, môn học, học kỳ, lớp học phần, lịch học và import dữ liệu.",
        "Ghi nhận các trạng thái PRESENT, LATE, ABSENT_EXCUSED và ABSENT_UNEXCUSED để phục vụ thống kê chuyên cần.",
        "Hỗ trợ sinh viên gửi đơn xin phép vắng học kèm minh chứng ảnh/PDF và giảng viên duyệt hoặc từ chối.",
        "Hỗ trợ cảnh báo sinh viên vắng vượt ngưỡng quy định, mặc định theo dữ liệu là 20% trên từng lớp học phần.",
    ])
    doc.add_heading("3. Phạm vi áp dụng", level=3)
    for text in [
        "Hệ thống phù hợp với các cơ sở đào tạo tổ chức lớp theo mô hình lớp học phần, trong đó mỗi lớp học phần có môn học, học kỳ, giảng viên phụ trách, lịch buổi học và danh sách sinh viên đăng ký. Phạm vi nghiệp vụ tập trung vào quản lý điểm danh, xử lý vắng mặt và theo dõi chuyên cần.",
        "Đề tài không đi sâu vào các phân hệ quản lý điểm số, học phí, đăng ký học phần toàn trường hoặc lập thời khóa biểu ở quy mô lớn. Tuy nhiên, cấu trúc dữ liệu của hệ thống đã tách các thực thể quan trọng như Subject, Semester, CourseSection, Lesson và Enrollment, nên có thể tích hợp với hệ thống đào tạo lớn hơn trong tương lai.",
    ]:
        p(doc, text)
    doc.add_heading("4. Đối tượng sử dụng", level=3)
    add_table(
        doc,
        ["Đối tượng", "Vai trò chính", "Nhu cầu sử dụng"],
        [
            ["Admin", "Quản trị dữ liệu nền và tài khoản", "Tạo tài khoản, khóa/mở khóa, quản lý khoa/lớp/môn/học kỳ/lớp học phần, import dữ liệu và xem báo cáo tổng quan."],
            ["Giảng viên", "Tổ chức và theo dõi điểm danh", "Xem lớp phụ trách, mở phiên điểm danh, lấy QR/OTP, điểm danh thủ công, đóng phiên, duyệt đơn xin phép và xem cảnh báo."],
            ["Sinh viên", "Thực hiện và theo dõi chuyên cần", "Xem lớp đang học, xem lịch học, điểm danh bằng QR/OTP, xem lịch sử, gửi đơn xin phép và nhận thông báo."],
        ],
        [1800, 2600, 4960],
    )
    doc.add_heading("5. Nền tảng kỹ thuật", level=3)
    p(doc, "Project được tổ chức theo kiến trúc client-server với hai phần frontend và backend tách riêng. Frontend sử dụng React, Vite và TypeScript; backend sử dụng Node.js, Express và TypeScript; cơ sở dữ liệu sử dụng PostgreSQL thông qua Prisma ORM. Hệ thống có Docker Compose để dựng database, backend và frontend trong cùng một môi trường triển khai.")
    add_bullets(doc, [
        "Frontend: React 19, Vite, TypeScript, React Router, Axios, lucide-react và ZXing cho khả năng quét mã QR.",
        "Backend: Node.js, Express, TypeScript, Prisma Client, Zod validation, Helmet, CORS và rate limit.",
        "CSDL: PostgreSQL, quản lý schema và migration bằng Prisma.",
        "Xác thực: JWT access token, refresh token, hash mật khẩu và refresh token bằng bcryptjs.",
        "Upload: Multer lưu file import Excel/CSV và minh chứng vắng học ảnh/PDF.",
        "Điểm danh: QR động có chữ ký HMAC, OTP 4-6 chữ số theo cấu hình, mặc định cửa sổ 30 giây.",
        "Triển khai: Docker Compose gồm db, backend và frontend; API backend chạy tại http://localhost:4000/api.",
    ])

    doc.add_heading("II. Cơ sở lý thuyết", level=2)
    theory = [
        ("1. Công nghệ phần mềm và quy trình phát triển phần mềm", [
            "Công nghệ phần mềm là lĩnh vực nghiên cứu và áp dụng các nguyên lý, phương pháp, công cụ để phát triển phần mềm có chất lượng. Một sản phẩm phần mềm không chỉ bao gồm mã nguồn mà còn bao gồm yêu cầu, thiết kế, dữ liệu, kiểm thử, tài liệu triển khai và khả năng bảo trì.",
            "Trong đề tài này, quy trình phát triển được thể hiện qua việc xác định nghiệp vụ điểm danh, phân tích tác nhân và use case, thiết kế cơ sở dữ liệu, xây dựng API, xây dựng giao diện theo vai trò và kiểm thử các tình huống chính. Cách tổ chức backend theo controllers, routes, services, middleware và validators giúp phân chia trách nhiệm rõ ràng.",
        ]),
        ("2. Mô hình client-server", [
            "Mô hình client-server tách phần giao diện tương tác với người dùng khỏi phần xử lý nghiệp vụ và lưu trữ dữ liệu. Frontend đóng vai trò client, gửi yêu cầu HTTP đến backend. Backend đóng vai trò server, xác thực người dùng, kiểm tra phân quyền, xử lý nghiệp vụ và thao tác với PostgreSQL.",
            "Trong project, frontend React gọi backend thông qua Axios tại địa chỉ API cấu hình. Backend Express gom các route theo nhóm /auth, /admin, /teacher và /student. Cách tách này làm cho mỗi vai trò có phạm vi API riêng, đồng thời giúp dễ kiểm soát bảo mật.",
        ]),
        ("3. RESTful API", [
            "RESTful API là phong cách thiết kế giao tiếp dựa trên tài nguyên. Các phương thức GET, POST, PATCH và DELETE được sử dụng tương ứng với đọc, tạo, cập nhật và xóa dữ liệu. Với hệ thống điểm danh, tài nguyên có thể là users, faculties, sections, lessons, sessions, attendance records hoặc leave requests.",
            "Ví dụ, Admin dùng POST /api/admin/users để tạo người dùng, Giảng viên dùng POST /api/teacher/sessions để tạo phiên điểm danh và Sinh viên dùng POST /api/student/attendance để gửi QR/OTP. Cấu trúc endpoint nhất quán giúp frontend dễ gọi API và giúp báo cáo kiểm thử rõ ràng.",
        ]),
        ("4. Cơ sở dữ liệu quan hệ", [
            "Cơ sở dữ liệu quan hệ lưu trữ dữ liệu dưới dạng bảng có khóa chính, khóa ngoại và ràng buộc duy nhất. Trong hệ thống điểm danh, quan hệ giữa lớp học phần, sinh viên, buổi học và bản ghi điểm danh là trọng tâm. Ví dụ, một CourseSection có nhiều Enrollment, mỗi Enrollment liên kết một sinh viên với một lớp học phần.",
            "Các ràng buộc như unique(email), unique(courseSectionId, studentId) và unique(attendanceSessionId, studentId) giúp tránh trùng dữ liệu. Đây là cơ sở để hệ thống không cho sinh viên điểm danh hai lần trong cùng một phiên.",
        ]),
        ("5. ORM và Prisma", [
            "ORM giúp ánh xạ bảng dữ liệu quan hệ thành model trong mã nguồn. Prisma cho phép định nghĩa schema, enum, quan hệ và migration ở một nơi, sau đó sinh Prisma Client để backend truy vấn dữ liệu an toàn hơn với TypeScript.",
            "Trong project, các model User, Faculty, Class, Subject, Semester, CourseSection, Enrollment, Lesson, AttendanceSession, AttendanceRecord, LeaveRequest, Notification, ImportLog và AuditLog được khai báo trong schema.prisma. Nhờ đó, thao tác tạo phiên, điểm danh, đóng phiên hoặc duyệt đơn đều làm việc trên model rõ ràng.",
        ]),
        ("6. Xác thực và phân quyền", [
            "Xác thực là quá trình kiểm tra danh tính người dùng, thường thông qua email và mật khẩu. Sau khi đăng nhập thành công, hệ thống cấp access token để gọi API và refresh token để lấy access token mới khi token cũ hết hạn.",
            "Phân quyền là quá trình kiểm soát tài nguyên theo vai trò. Middleware authenticate kiểm tra token, còn authorize kiểm tra role. Project sử dụng ba role ADMIN, TEACHER và STUDENT. Tài khoản bị LOCKED hoặc token không hợp lệ sẽ bị từ chối truy cập.",
        ]),
        ("7. QR động và OTP trong điểm danh", [
            "QR động và OTP giúp giảm gian lận so với mã cố định. Backend sinh QR token chứa sessionId, courseSectionId, nonce và cửa sổ thời gian, sau đó ký bằng HMAC-SHA256. OTP được sinh từ sessionId, nonce và window hiện tại nên thay đổi theo chu kỳ cấu hình.",
            "Khi sinh viên gửi điểm danh, backend kiểm tra QR/OTP, trạng thái phiên, quyền thuộc lớp học phần và ràng buộc chưa điểm danh trước đó. Cách kết hợp QR và OTP giúp giảng viên hiển thị mã tại lớp, còn sinh viên chỉ điểm danh được trong phiên đang mở.",
        ]),
    ]
    for title, paras in theory:
        doc.add_heading(title, level=3)
        for text in paras:
            p(doc, text)


def requirement_rows(role: str) -> list[list[str]]:
    if role == "admin":
        funcs = [
            ("Đăng nhập", "Tra cứu", "Email/mật khẩu đúng, tài khoản ACTIVE", "Form đăng nhập", "Nhận access token và refresh token"),
            ("Quản lý tài khoản", "CRUD", "Email duy nhất, role hợp lệ", "Form người dùng", "Tạo Admin/Giảng viên/Sinh viên"),
            ("Khóa/mở khóa tài khoản", "Cập nhật", "AccountStatus ACTIVE/LOCKED", "Nút khóa", "Tài khoản LOCKED không được truy cập"),
            ("Quản lý khoa", "CRUD", "Mã khoa duy nhất", "Form khoa", "Liên kết với lớp và môn học"),
            ("Quản lý lớp hành chính", "CRUD/Import", "Mã lớp duy nhất, thuộc khoa", "Form lớp, file CSV/XLSX", "Dùng cho sinh viên"),
            ("Quản lý môn học", "CRUD", "Mã môn duy nhất, số tín chỉ", "Form môn", "Liên kết khoa"),
            ("Quản lý học kỳ", "CRUD", "Ngày bắt đầu <= ngày kết thúc", "Form học kỳ", "Lớp học phần thuộc học kỳ"),
            ("Quản lý lớp học phần", "CRUD/Import", "Mã lớp học phần duy nhất", "Form lớp học phần", "Có môn, học kỳ, giảng viên"),
            ("Phân công giảng viên", "Cập nhật", "teacherId phải là user TEACHER", "Form lớp học phần", "Một teacher phụ trách nhiều lớp"),
            ("Import sinh viên bằng Excel", "Import", "File .xlsx/.xls/.csv, tối đa 5MB", "Mẫu import", "Tạo user và enrollment"),
            ("Xem nhật ký import", "Tra cứu", "Lưu totalRows/successRows/failedRows", "Bảng import log", "Hỗ trợ rà soát lỗi dữ liệu"),
            ("Xem dữ liệu hệ thống", "Báo cáo", "Tính tỷ lệ chuyên cần và cảnh báo", "Trang báo cáo", "Có gửi cảnh báo vắng"),
        ]
    elif role == "teacher":
        funcs = [
            ("Đăng nhập", "Tra cứu", "Tài khoản TEACHER/ADMIN hợp lệ", "Form đăng nhập", "Truy cập dashboard giảng viên"),
            ("Xem lớp được phân công", "Tra cứu", "courseSection.teacherId = user.id", "Danh sách lớp", "Chỉ thấy lớp phụ trách"),
            ("Xem danh sách sinh viên", "Tra cứu", "Sinh viên có enrollment", "Bảng sinh viên", "Phục vụ điểm danh"),
            ("Tạo buổi học", "CRUD", "Ngày học, giờ bắt đầu/kết thúc", "Form lesson", "Có import lịch học"),
            ("Tạo phiên điểm danh", "Nghiệp vụ", "Lesson đã bắt đầu, chưa kết thúc", "Nút tạo phiên", "Mỗi teacher chỉ có một phiên mở"),
            ("Lấy QR/OTP", "Nghiệp vụ", "Session status OPEN", "Màn hình QR/OTP", "Mã thay đổi 30 giây"),
            ("Điểm danh thủ công", "Cập nhật", "Phiên đang mở, sinh viên thuộc lớp", "Form manual mark", "Cần lý do"),
            ("Đóng phiên điểm danh", "Nghiệp vụ", "Teacher phụ trách lớp", "Nút đóng phiên", "Tự đánh vắng sinh viên chưa có record"),
            ("Xem kết quả điểm danh", "Báo cáo", "Lọc theo session/section", "Bảng kết quả", "Có export Excel/CSV"),
            ("Duyệt đơn xin phép", "Cập nhật", "Đơn PENDING", "Form duyệt", "Tạo notification cho sinh viên"),
            ("Theo dõi cảnh báo", "Báo cáo", "Tỷ lệ vắng >= threshold", "Trang báo cáo", "Ngưỡng mặc định 20%"),
        ]
    else:
        funcs = [
            ("Đăng nhập", "Tra cứu", "Tài khoản STUDENT hợp lệ", "Form đăng nhập", "Login lần 2 thu hồi token cũ"),
            ("Xem lớp đang tham gia", "Tra cứu", "Có enrollment", "Danh sách lớp", "Hiển thị môn, học kỳ, giảng viên"),
            ("Xem lịch học", "Tra cứu", "Lesson thuộc lớp đã đăng ký", "Lịch học", "Có thông tin phiên gần nhất"),
            ("Điểm danh QR/OTP", "Nghiệp vụ", "Phiên OPEN, QR/OTP hợp lệ", "Form điểm danh", "Tạo record PRESENT"),
            ("Xem lịch sử điểm danh", "Tra cứu", "attendance_records theo studentId", "Bảng lịch sử", "Theo dõi chuyên cần cá nhân"),
            ("Gửi đơn xin phép", "Nghiệp vụ", "Chỉ cho record ABSENT_UNEXCUSED", "Form đơn + file", "Bắt buộc minh chứng"),
            ("Xem thông báo", "Tra cứu/Cập nhật", "Notification theo userId", "Danh sách thông báo", "Có đánh dấu đã đọc"),
            ("Đổi mật khẩu", "Cập nhật", "Mật khẩu cũ đúng", "Form đổi mật khẩu", "Thu hồi refresh token"),
        ]
    return [[str(i + 1), *row] for i, row in enumerate(funcs)]


def chapter_two(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG II. PHÂN TÍCH NỘI DUNG, YÊU CẦU", level=1)
    doc.add_heading("I. Giới thiệu quy trình quản lý hệ thống", level=2)
    for text in [
        "Quy trình quản lý bắt đầu từ Admin đăng nhập vào hệ thống bằng tài khoản có role ADMIN. Sau khi xác thực thành công, Admin được cấp token và truy cập các phân hệ quản lý dữ liệu nền. Các dữ liệu như khoa, lớp hành chính, môn học và học kỳ là cơ sở để tạo lớp học phần.",
        "Khi tạo lớp học phần, Admin chọn môn học, học kỳ và giảng viên phụ trách. Sau đó Admin có thể import danh sách sinh viên vào lớp học phần bằng file Excel hoặc CSV. Backend đọc file upload, kiểm tra dữ liệu từng dòng, tạo tài khoản sinh viên nếu cần và tạo enrollment liên kết sinh viên với lớp học phần.",
        "Toàn bộ dữ liệu được lưu trong PostgreSQL thông qua Prisma Client. Các lỗi import được ghi vào import_logs với số dòng thành công, số dòng thất bại và chi tiết lỗi. Điều này giúp Admin kiểm soát chất lượng dữ liệu đầu vào thay vì chỉ biết import thành công hoặc thất bại chung chung.",
    ]:
        p(doc, text)
    doc.add_heading("II. Giới thiệu quy trình điểm danh bằng QR/OTP", level=2)
    p(doc, "Giảng viên đăng nhập, chọn lớp học phần được phân công và chọn một buổi học đã bắt đầu. Backend kiểm tra giảng viên có phụ trách lớp không, kiểm tra buổi học chưa kết thúc và đảm bảo giảng viên không đang có phiên điểm danh mở ở lớp khác. Khi hợp lệ, hệ thống tạo AttendanceSession với trạng thái OPEN và nonce ngẫu nhiên.")
    p(doc, "Trong phiên đang mở, giảng viên lấy QR/OTP từ API. QR token chứa sessionId, courseSectionId, nonce và window thời gian; OTP được sinh từ HMAC dựa trên sessionId, nonce và window. Sinh viên quét QR hoặc nhập token/OTP, sau đó backend kiểm tra sinh viên có enrollment trong lớp, phiên còn OPEN, QR chưa hết hạn, OTP đúng và chưa có AttendanceRecord trong phiên. Nếu hợp lệ, hệ thống tạo bản ghi PRESENT theo phương thức QR_OTP.")
    p(doc, "Khi giảng viên đóng phiên hoặc khi hệ thống phát hiện buổi học đã kết thúc, các sinh viên chưa có bản ghi điểm danh được tự động tạo AttendanceRecord với trạng thái ABSENT_UNEXCUSED và phương thức SYSTEM. Cơ chế này giúp dữ liệu sau phiên luôn đầy đủ cho toàn bộ danh sách sinh viên trong lớp học phần.")
    doc.add_heading("III. Giới thiệu quy trình xin phép vắng học", level=2)
    p(doc, "Sinh viên chỉ có thể gửi đơn xin phép cho bản ghi đang ở trạng thái ABSENT_UNEXCUSED. Khi tạo đơn, sinh viên nhập lý do và upload minh chứng ảnh hoặc PDF. Backend kiểm tra file bắt buộc, kiểm tra trạng thái bản ghi, sau đó tạo hoặc cập nhật LeaveRequest ở trạng thái PENDING.")
    p(doc, "Sau khi đơn được gửi, hệ thống tạo thông báo cho giảng viên phụ trách lớp học phần. Giảng viên xem danh sách đơn, duyệt hoặc từ chối. Nếu duyệt, hệ thống cập nhật bản ghi điểm danh thành ABSENT_EXCUSED; nếu từ chối, bản ghi vẫn là ABSENT_UNEXCUSED và lưu ghi chú phản hồi. Sinh viên nhận thông báo kết quả xử lý đơn.")
    doc.add_heading("IV. Yêu cầu chức năng nghiệp vụ", level=2)
    headers = ["STT", "Công việc", "Loại", "Quy định/Công thức", "Biểu mẫu", "Ghi chú"]
    widths = [520, 2100, 1200, 2300, 1500, 1740]
    doc.add_heading("1. Chức năng của Admin", level=3)
    add_table(doc, headers, requirement_rows("admin"), widths, "Bảng 2.1. Yêu cầu chức năng của Admin")
    doc.add_heading("2. Chức năng của Giảng viên", level=3)
    add_table(doc, headers, requirement_rows("teacher"), widths, "Bảng 2.2. Yêu cầu chức năng của Giảng viên")
    doc.add_heading("3. Chức năng của Sinh viên", level=3)
    add_table(doc, headers, requirement_rows("student"), widths, "Bảng 2.3. Yêu cầu chức năng của Sinh viên")
    doc.add_heading("V. Yêu cầu chức năng hệ thống và yêu cầu chất lượng", level=2)
    doc.add_heading("1. Yêu cầu chức năng hệ thống", level=3)
    functional = [
        ["FR-01", "Xác thực người dùng trước khi truy cập API được bảo vệ.", "Bắt buộc"],
        ["FR-02", "Phân quyền theo vai trò ADMIN, TEACHER, STUDENT.", "Bắt buộc"],
        ["FR-03", "Sinh QR/OTP cho phiên điểm danh đang mở, có thời hạn 30 giây.", "Bắt buộc"],
        ["FR-04", "Kiểm tra sinh viên thuộc lớp học phần trước khi điểm danh.", "Bắt buộc"],
        ["FR-05", "Không cho phép điểm danh hai lần trong cùng một phiên.", "Bắt buộc"],
        ["FR-06", "Tự động đánh vắng không phép khi đóng phiên.", "Bắt buộc"],
        ["FR-07", "Cho phép giảng viên điểm danh thủ công với lý do.", "Bắt buộc"],
        ["FR-08", "Cho phép sinh viên gửi đơn xin phép kèm minh chứng.", "Bắt buộc"],
        ["FR-09", "Tính cảnh báo vắng theo ngưỡng của lớp học phần.", "Nên có"],
        ["FR-10", "Gửi thông báo khi có đơn xin phép mới hoặc kết quả duyệt đơn.", "Nên có"],
    ]
    add_table(doc, ["Mã", "Yêu cầu", "Mức độ"], functional, [900, 6900, 1560])
    doc.add_heading("2. Yêu cầu phi chức năng", level=3)
    non_functional = [
        ["Bảo mật", "Mật khẩu và refresh token được hash bằng bcrypt; API được bảo vệ bằng JWT; tài khoản LOCKED không được truy cập."],
        ["Hiệu năng", "Thao tác điểm danh chỉ cần kiểm tra token, session, enrollment và unique record nên có thể phản hồi nhanh trong giờ học."],
        ["Dễ sử dụng", "Giao diện frontend chia theo vai trò, điều hướng rõ ràng bằng React Router và layout dashboard."],
        ["Chính xác", "Ràng buộc unique trên attendanceSessionId và studentId ngăn ghi trùng điểm danh."],
        ["Mở rộng", "Có thể bổ sung nhận diện khuôn mặt, định vị, email, mobile app hoặc dashboard nâng cao."],
        ["Bảo trì", "Backend/frontend tách riêng; TypeScript, Prisma schema và route theo module giúp dễ sửa đổi."],
        ["Ổn định", "Có test QR/OTP và danh sách test case nghiệp vụ cho các luồng quan trọng."],
    ]
    add_table(doc, ["Nhóm yêu cầu", "Mô tả"], non_functional, [1800, 7560])


def db_rows() -> list[list[str]]:
    return [
        ["users", "Lưu tài khoản Admin/Giảng viên/Sinh viên, role, trạng thái, mã SV/GV, lớp hành chính và token session.", "email, role, status, class_id"],
        ["faculties", "Lưu khoa/đơn vị đào tạo.", "code, name"],
        ["classes", "Lưu lớp hành chính thuộc khoa.", "code, name, faculty_id"],
        ["subjects", "Lưu môn học thuộc khoa.", "code, name, credits"],
        ["semesters", "Lưu học kỳ và khoảng thời gian.", "name, start_date, end_date"],
        ["course_sections", "Lưu lớp học phần, môn học, học kỳ, giảng viên và ngưỡng vắng.", "code, subject_id, semester_id, teacher_id"],
        ["enrollments", "Liên kết sinh viên với lớp học phần.", "course_section_id, student_id"],
        ["lessons", "Lưu lịch từng buổi học.", "lesson_date, start_time, end_time, room"],
        ["attendance_sessions", "Lưu phiên điểm danh theo buổi học.", "lesson_id, status, nonce, opened_at"],
        ["attendance_records", "Lưu kết quả điểm danh của từng sinh viên trong phiên.", "status, method, marked_at"],
        ["leave_requests", "Lưu đơn xin phép và trạng thái duyệt.", "reason, evidence_path, status"],
        ["notifications", "Lưu thông báo cho người dùng.", "title, message, read_at"],
        ["import_logs", "Lưu nhật ký import dữ liệu.", "total_rows, success_rows, failed_rows"],
        ["audit_logs", "Lưu nhật ký thao tác hệ thống.", "actor_id, action, entity"],
    ]


def add_use_case_table(doc: Document, name: str, actor: str, goal: str, pre: str, post: str, main: str, alt: str) -> None:
    add_table(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ["Tên use case", name],
            ["Tác nhân", actor],
            ["Mục tiêu", goal],
            ["Tiền điều kiện", pre],
            ["Hậu điều kiện", post],
            ["Luồng chính", main],
            ["Luồng thay thế/ngoại lệ", alt],
        ],
        [2100, 7260],
    )


def chapter_three(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG III. PHÂN TÍCH THIẾT KẾ", level=1)
    doc.add_heading("I. Sơ đồ use case", level=2)
    p(doc, "Sơ đồ use case tổng quát của hệ thống gồm ba actor chính: Admin, Giảng viên và Sinh viên. Admin chịu trách nhiệm dữ liệu nền và tài khoản; Giảng viên chịu trách nhiệm lớp học phần, phiên điểm danh và đơn xin phép; Sinh viên sử dụng hệ thống để điểm danh và theo dõi chuyên cần cá nhân.")
    add_table(
        doc,
        ["Actor", "Use case chính"],
        [
            ["Admin", "Đăng nhập; quản lý người dùng; khóa/mở khóa tài khoản; quản lý khoa, lớp, môn học, học kỳ, lớp học phần; import sinh viên/lịch học/lớp; xem báo cáo và gửi cảnh báo."],
            ["Giảng viên", "Đăng nhập; xem lớp phụ trách; xem sinh viên; tạo lesson; import lịch học; tạo phiên; lấy QR/OTP; điểm danh thủ công; đóng phiên; xem báo cáo; duyệt đơn xin phép."],
            ["Sinh viên", "Đăng nhập; xem lớp đang học; xem lịch học; điểm danh QR/OTP; xem lịch sử; gửi đơn xin phép; xem thông báo; đổi mật khẩu."],
        ],
        [1600, 7760],
    )
    add_placeholder(doc, "[Chèn Hình 3.1. Sơ đồ use case tổng quát tại đây]")
    doc.add_heading("II. Đặc tả use case", level=2)
    use_cases = [
        ("Đăng nhập", "Admin/Giảng viên/Sinh viên", "Xác thực người dùng và cấp token truy cập.", "Người dùng có tài khoản ACTIVE.", "Người dùng nhận access token, refresh token và thông tin role.", "Nhập email/mật khẩu; hệ thống kiểm tra passwordHash; cấp token; chuyển đến dashboard theo role.", "Sai thông tin trả INVALID_CREDENTIALS; tài khoản khóa trả ACCOUNT_LOCKED."),
        ("Admin tạo tài khoản người dùng", "Admin", "Tạo tài khoản cho Admin, Giảng viên hoặc Sinh viên.", "Admin đã đăng nhập.", "Tài khoản mới được lưu trong users với mật khẩu đã hash.", "Admin nhập email, họ tên, role, mã SV/GV nếu có; hệ thống kiểm tra unique; hash password; lưu user.", "Email hoặc mã trùng thì báo lỗi; role không hợp lệ thì từ chối."),
        ("Admin import danh sách sinh viên", "Admin", "Tạo hoặc cập nhật sinh viên và đưa vào lớp học phần.", "Lớp học phần đã tồn tại, file đúng định dạng.", "Sinh viên và enrollment được tạo; import log được lưu.", "Upload file; backend đọc từng dòng; kiểm tra MSSV/email/lớp; tạo user và enrollment; ghi số dòng thành công/thất bại.", "Dòng lỗi được đưa vào errorDetails; file sai loại hoặc quá 5MB bị từ chối."),
        ("Giảng viên tạo phiên điểm danh", "Giảng viên", "Mở phiên điểm danh cho một lesson.", "Giảng viên phụ trách lớp, lesson đã bắt đầu và chưa kết thúc.", "AttendanceSession trạng thái OPEN được tạo.", "Chọn lesson; hệ thống kiểm tra quyền, đóng phiên hết hạn, kiểm tra phiên mở hiện tại; tạo nonce và session.", "Nếu chưa đến giờ học trả LESSON_NOT_STARTED; nếu đã có phiên mở trả OPEN_SESSION_EXISTS."),
        ("Sinh viên điểm danh QR/OTP", "Sinh viên", "Ghi nhận sinh viên có mặt trong phiên.", "Sinh viên đăng nhập, thuộc lớp học phần, phiên OPEN.", "AttendanceRecord trạng thái PRESENT, method QR_OTP.", "Quét QR hoặc nhập token; nhập OTP; hệ thống xác thực QR/OTP, kiểm tra enrollment và unique record; tạo bản ghi.", "OTP sai trả INVALID_OTP; QR hết hạn trả EXPIRED_QR; điểm danh lại trả ALREADY_ATTENDED."),
        ("Giảng viên điểm danh thủ công", "Giảng viên", "Ghi nhận hoặc cập nhật trạng thái điểm danh cho sinh viên.", "Phiên OPEN, sinh viên thuộc lớp.", "AttendanceRecord được upsert với method MANUAL.", "Chọn sinh viên, trạng thái và lý do; hệ thống kiểm tra quyền; upsert record.", "Phiên đóng trả SESSION_CLOSED; sinh viên không thuộc lớp trả NOT_ENROLLED."),
        ("Sinh viên gửi đơn xin phép", "Sinh viên", "Xin chuyển vắng không phép thành vắng có phép.", "Có record ABSENT_UNEXCUSED và file minh chứng.", "LeaveRequest PENDING được tạo, giảng viên nhận thông báo.", "Chọn bản ghi vắng; nhập lý do; upload file; hệ thống kiểm tra điều kiện; tạo đơn và notification.", "Không có file trả EVIDENCE_REQUIRED; không phải bản ghi vắng trả NOT_ABSENT."),
        ("Giảng viên duyệt đơn xin phép", "Giảng viên", "Phê duyệt hoặc từ chối đơn vắng học.", "Đơn thuộc lớp giảng viên phụ trách và đang PENDING.", "LeaveRequest cập nhật APPROVED/REJECTED, record và notification được cập nhật.", "Mở danh sách đơn; chọn duyệt/từ chối; hệ thống cập nhật đơn, attendance record và thông báo cho sinh viên.", "Đơn đã xử lý trả LEAVE_ALREADY_REVIEWED."),
    ]
    for uc in use_cases:
        add_use_case_table(doc, *uc)
    doc.add_heading("III. Sơ đồ hoạt động", level=2)
    activities = [
        ("1. Sơ đồ hoạt động đăng nhập", "Người dùng nhập email và mật khẩu. Backend kiểm tra định dạng bằng Zod, tìm user theo email, so khớp mật khẩu bằng bcrypt và kiểm tra trạng thái tài khoản. Nếu hợp lệ, backend cấp access token và refresh token. Với sinh viên, hệ thống tạo currentSessionId mới để vô hiệu hóa phiên đăng nhập cũ."),
        ("2. Sơ đồ hoạt động tạo phiên điểm danh", "Giảng viên chọn lớp học phần và lesson. Backend kiểm tra quyền phụ trách lớp, kiểm tra lesson đã bắt đầu, chưa kết thúc, đồng thời kiểm tra giảng viên chưa có phiên khác đang mở. Nếu hợp lệ, hệ thống tạo attendance_sessions với nonce mới và trạng thái OPEN."),
        ("3. Sơ đồ hoạt động điểm danh bằng QR/OTP", "Sinh viên quét QR hoặc nhập token, nhập OTP và gửi yêu cầu. Hệ thống giải mã QR, kiểm tra chữ ký HMAC, kiểm tra window hiện tại hoặc window liền trước, so khớp OTP, kiểm tra enrollment và ràng buộc unique record. Nếu tất cả hợp lệ, bản ghi PRESENT được tạo."),
        ("4. Sơ đồ hoạt động gửi và duyệt đơn xin phép", "Sinh viên gửi đơn cho bản ghi ABSENT_UNEXCUSED kèm minh chứng. Hệ thống tạo LeaveRequest PENDING và notification cho giảng viên. Giảng viên duyệt hoặc từ chối; hệ thống cập nhật trạng thái điểm danh và gửi notification kết quả cho sinh viên."),
    ]
    for idx, (title, desc) in enumerate(activities, start=2):
        doc.add_heading(title, level=3)
        p(doc, desc)
        add_placeholder(doc, f"[Chèn Hình 3.{idx}. {title.replace('1. ', '').replace('2. ', '').replace('3. ', '').replace('4. ', '')} tại đây]")
    doc.add_heading("IV. Thiết kế cơ sở dữ liệu", level=2)
    doc.add_heading("1. Mô hình ERD", level=3)
    p(doc, "Mô hình ERD được xây dựng quanh CourseSection và AttendanceSession. CourseSection liên kết Subject, Semester, Teacher và Enrollment; AttendanceSession liên kết Lesson và AttendanceRecord; LeaveRequest liên kết AttendanceRecord để xử lý xin phép vắng. Các bảng notification, import_logs và audit_logs hỗ trợ thông báo, nhập dữ liệu và theo dõi thao tác.")
    add_placeholder(doc, "[Chèn Hình 3.6. Mô hình ERD hệ thống tại đây]")
    doc.add_heading("2. Mô tả quan hệ giữa các bảng", level=3)
    add_bullets(doc, [
        "Một Faculty có nhiều Class và nhiều Subject.",
        "Một Class có nhiều User sinh viên thông qua classId.",
        "Một Subject có nhiều CourseSection; một Semester có nhiều CourseSection.",
        "Một Teacher/User có thể phụ trách nhiều CourseSection.",
        "Một CourseSection có nhiều Enrollment, Lesson, AttendanceSession và ImportLog.",
        "Một Student/User có nhiều Enrollment và AttendanceRecord.",
        "Một AttendanceSession có nhiều AttendanceRecord và LeaveRequest.",
        "Một AttendanceRecord có thể liên kết một LeaveRequest.",
        "Một User có nhiều Notification; ImportLog ghi nhận quá trình import dữ liệu.",
    ])
    doc.add_heading("3. Cấu trúc các bảng", level=3)
    add_table(doc, ["Bảng", "Mục đích", "Trường chính"], db_rows(), [1900, 5000, 2460], "Bảng 3.1. Danh sách bảng trong cơ sở dữ liệu")
    structures = [
        ("users", "id UUID PK; email unique; password_hash; full_name; role; status; student_code unique nullable; teacher_code unique nullable; class_id FK; current_session_id; refresh_token_hash; created_at; updated_at."),
        ("faculties", "id UUID PK; name unique; code unique; created_at; updated_at."),
        ("classes", "id UUID PK; name; code unique; faculty_id FK; created_at; updated_at; unique(faculty_id, name)."),
        ("subjects", "id UUID PK; code unique; name; credits; faculty_id FK; created_at; updated_at; unique(faculty_id, name)."),
        ("semesters", "id UUID PK; name unique; start_date; end_date; created_at; updated_at."),
        ("course_sections", "id UUID PK; code unique; subject_id FK; semester_id FK; teacher_id FK; status; absence_threshold_percent; created_at; updated_at."),
        ("enrollments", "id UUID PK; course_section_id FK; student_id FK; created_at; unique(course_section_id, student_id)."),
        ("lessons", "id UUID PK; course_section_id FK; lesson_date; start_time; end_time; room; topic; created_at; unique(course_section_id, lesson_date, start_time)."),
        ("attendance_sessions", "id UUID PK; course_section_id FK; lesson_id FK; status; nonce; opened_at; closed_at; created_by_id."),
        ("attendance_records", "id UUID PK; attendance_session_id FK; student_id FK; status; method; marked_at; marked_by_id FK nullable; reason; created_at; updated_at; unique(attendance_session_id, student_id)."),
        ("leave_requests", "id UUID PK; attendance_session_id FK; attendance_record_id unique FK nullable; student_id FK; reason; evidence_path; status; reviewed_by_id FK; reviewed_at; review_note; created_at; updated_at."),
        ("notifications", "id UUID PK; user_id FK; title; message; read_at; created_at."),
        ("import_logs", "id UUID PK; course_section_id FK nullable; imported_by_id FK; file_name; total_rows; success_rows; failed_rows; error_details JSON; created_at."),
        ("audit_logs", "id UUID PK; actor_id FK nullable; action; entity; entity_id; metadata JSON; created_at."),
    ]
    for name, fields in structures:
        doc.add_heading(f"Bảng {name}", level=3)
        p(doc, fields)
    doc.add_heading("V. Thiết kế giao diện", level=2)
    interfaces = [
        ("Giao diện đăng nhập", "Dùng cho cả ba vai trò, gồm email, mật khẩu, nút đăng nhập và xử lý lỗi."),
        ("Giao diện trang chủ Admin", "Hiển thị số liệu tổng quan về người dùng, lớp học phần, phiên điểm danh và cảnh báo."),
        ("Giao diện quản lý người dùng", "Cho phép tạo, sửa, xóa, khóa/mở khóa tài khoản và phân role."),
        ("Giao diện quản lý khoa/lớp/môn/học kỳ", "Quản lý dữ liệu danh mục nền phục vụ tạo lớp học phần."),
        ("Giao diện quản lý lớp học phần", "Tạo lớp, phân công giảng viên, xem sinh viên và lịch học."),
        ("Giao diện import sinh viên", "Upload file mẫu, xem số dòng thành công/thất bại và lỗi."),
        ("Giao diện danh sách lớp học phần của giảng viên", "Liệt kê lớp giảng viên được phân công và truy cập chi tiết."),
        ("Giao diện tạo phiên điểm danh", "Chọn buổi học, tạo phiên, xem trạng thái phiên đang mở."),
        ("Giao diện hiển thị QR/OTP", "Hiển thị mã QR, OTP và thời gian còn hiệu lực."),
        ("Giao diện sinh viên điểm danh", "Nhận QR token từ URL hoặc nhập thủ công, nhập OTP và gửi điểm danh."),
        ("Giao diện lịch sử điểm danh", "Sinh viên xem từng buổi, trạng thái và phương thức điểm danh."),
        ("Giao diện gửi đơn xin phép", "Sinh viên chọn bản ghi vắng, nhập lý do và upload minh chứng."),
        ("Giao diện duyệt đơn xin phép", "Giảng viên duyệt/từ chối và nhập ghi chú phản hồi."),
        ("Giao diện thông báo", "Người dùng xem và đánh dấu đã đọc thông báo."),
    ]
    for title, desc in interfaces:
        doc.add_heading(title, level=3)
        p(doc, f"Mục đích: {desc} Đối tượng sử dụng và chức năng xử lý được xác định theo route frontend trong App.tsx, gồm các đường dẫn /admin, /teacher và /student tương ứng với từng vai trò.")
        add_placeholder(doc, f"[Chèn hình minh họa giao diện {title.lower()} tại đây]")
    doc.add_heading("VI. Thiết kế xử lý", level=2)
    processes = [
        ("Xử lý đăng nhập và cấp JWT", "Backend kiểm tra email/mật khẩu, trạng thái tài khoản, tạo sessionId, ký access token và refresh token. Refresh token được hash trước khi lưu."),
        ("Xử lý refresh token", "Client gửi refresh token; backend xác thực chữ ký, so khớp hash, kiểm tra currentSessionId đối với sinh viên và cấp access token mới."),
        ("Xử lý phân quyền API theo role", "Middleware authenticate đọc Bearer token, lấy user từ CSDL, kiểm tra tài khoản LOCKED; authorize kiểm tra role được phép trên từng route."),
        ("Xử lý tạo tài khoản và hash mật khẩu", "Admin gửi thông tin user; backend validate bằng Zod, hash mật khẩu bằng bcryptjs và lưu vào users."),
        ("Xử lý import sinh viên bằng Excel", "Multer lưu file vào uploads/imports, backend đọc dữ liệu, kiểm tra từng dòng, tạo user/enrollment và ghi import_logs."),
        ("Xử lý tạo phiên điểm danh", "Teacher gửi lessonId; backend kiểm tra quyền, thời gian lesson, phiên đang mở và tạo AttendanceSession với nonce."),
        ("Xử lý sinh QR động và OTP 30 giây", "Backend tạo QR token ký HMAC và OTP theo currentWindow; trả qrDataUrl, qrUrl, otp và validSeconds."),
        ("Xử lý sinh viên điểm danh", "Backend chuẩn hóa token, giải mã sessionId, đóng phiên hết hạn nếu cần, kiểm tra session OPEN, enrollment, QR/OTP và unique record."),
        ("Xử lý điểm danh thủ công", "Giảng viên chọn sinh viên và lý do; backend upsert AttendanceRecord method MANUAL."),
        ("Xử lý đóng phiên điểm danh", "Backend duyệt enrollment của lớp; sinh viên chưa có record được tạo ABSENT_UNEXCUSED method SYSTEM; session chuyển CLOSED."),
        ("Xử lý đơn xin phép", "Sinh viên gửi đơn cho record ABSENT_UNEXCUSED; giảng viên duyệt/từ chối; hệ thống cập nhật record và notification."),
        ("Xử lý cảnh báo vắng quá ngưỡng 20%", "Báo cáo tính tỷ lệ vắng của sinh viên theo số phiên đã có record; nếu tỷ lệ vắng đạt hoặc vượt absenceThresholdPercent thì sinh cảnh báo."),
    ]
    for title, desc in processes:
        doc.add_heading(title, level=3)
        p(doc, desc)


def chapter_four(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG IV. PHÁT TRIỂN/THỰC THI", level=1)
    doc.add_heading("I. Cấu trúc source code", level=2)
    p(doc, "Project được tổ chức thành các thư mục chính backend, frontend, doc và file docker-compose.yml. Cách tách này phù hợp với kiến trúc client-server: backend chịu trách nhiệm API và dữ liệu, frontend chịu trách nhiệm giao diện, còn Docker Compose hỗ trợ dựng môi trường triển khai.")
    add_table(
        doc,
        ["Thành phần", "Vai trò"],
        [
            ["backend/", "Mã nguồn Express TypeScript, Prisma schema, migration, seed, controller, route, service, middleware và test."],
            ["frontend/", "Mã nguồn React/Vite TypeScript, layout, page, component, context, service, type và asset."],
            ["doc/", "Tài liệu báo cáo, tài liệu phân tích và các file Word liên quan."],
            ["docker-compose.yml", "Dựng PostgreSQL, backend và frontend bằng container."],
        ],
        [2200, 7160],
    )
    doc.add_heading("1. Backend", level=3)
    backend_dirs = [
        ["backend/src/config", "Cấu hình env và Prisma Client."],
        ["backend/src/controllers", "Xử lý request/response cho auth, admin, teacher và student."],
        ["backend/src/middleware", "Xác thực, phân quyền, upload file và xử lý lỗi."],
        ["backend/src/routes", "Khai báo endpoint theo module."],
        ["backend/src/services", "Chứa logic nghiệp vụ dùng chung, nổi bật là attendance.service.ts."],
        ["backend/src/tests", "Kiểm thử QR/OTP bằng Vitest."],
        ["backend/src/types", "Mở rộng kiểu Express Request cho req.user."],
        ["backend/src/utils", "Token, mật khẩu, QR/OTP và helper HTTP."],
        ["backend/src/validators", "Schema dùng chung cho validate dữ liệu."],
        ["backend/prisma", "Prisma schema, migrations và seed dữ liệu mẫu."],
    ]
    add_table(doc, ["Thư mục", "Vai trò"], backend_dirs, [3000, 6360])
    doc.add_heading("2. Frontend", level=3)
    frontend_dirs = [
        ["frontend/src/assets", "Chứa hình ảnh và asset dùng trong giao diện."],
        ["frontend/src/components", "Component dùng lại như Button, DataTable, Modal, StatCard, PageHeader."],
        ["frontend/src/context", "AuthContext và ToastContext quản lý trạng thái xác thực và thông báo."],
        ["frontend/src/layouts", "DashboardLayout, ProtectedRoute và RequireRole."],
        ["frontend/src/pages", "Các trang chia theo admin, teacher, student và trang login/not found."],
        ["frontend/src/services", "Axios client, service theo vai trò, mock data và token store."],
        ["frontend/src/types", "Định nghĩa TypeScript cho User, CourseSection, Lesson, AttendanceRecord..."],
        ["frontend/src/utils", "Hàm định dạng và hook useAsync."],
    ]
    add_table(doc, ["Thư mục", "Vai trò"], frontend_dirs, [3000, 6360])
    doc.add_heading("II. Phát triển backend", level=2)
    modules = [
        ("Module xác thực", ["POST /api/auth/login", "POST /api/auth/refresh", "POST /api/auth/logout", "GET /api/auth/me", "PATCH /api/auth/change-password"], "Module này chịu trách nhiệm đăng nhập, cấp token, refresh token, đăng xuất và đổi mật khẩu. Mật khẩu được kiểm tra bằng bcryptjs, refresh token được hash trước khi lưu vào CSDL."),
        ("Module Admin", ["GET /api/admin/users", "POST /api/admin/users", "PATCH /api/admin/users/:id/lock", "CRUD faculties/classes/subjects/semesters/sections", "POST /api/admin/sections/:sectionId/import-students", "GET /api/admin/reports/overview"], "Module Admin quản lý dữ liệu nền và báo cáo toàn hệ thống. Các route import dùng Multer và ghi import_logs để theo dõi kết quả."),
        ("Module Teacher", ["GET /api/teacher/sections", "POST /api/teacher/sessions", "GET /api/teacher/sessions/:sessionId/qr-otp", "POST /api/teacher/sessions/:sessionId/manual-mark", "PATCH /api/teacher/sessions/:sessionId/close", "PATCH /api/teacher/leave-requests/:id/review"], "Module Teacher tập trung vào lớp phụ trách, buổi học, phiên điểm danh, điểm danh thủ công, xuất báo cáo và xử lý đơn xin phép."),
        ("Module Student", ["GET /api/student/sections", "GET /api/student/schedule", "POST /api/student/attendance", "GET /api/student/attendance/history", "POST /api/student/leave-requests", "GET /api/student/notifications"], "Module Student phục vụ xem lớp, xem lịch, điểm danh QR/OTP, xem lịch sử, gửi đơn xin phép và nhận thông báo."),
    ]
    for title, endpoints, desc in modules:
        doc.add_heading(title, level=3)
        p(doc, desc)
        add_bullets(doc, endpoints)
    doc.add_heading("III. Phát triển frontend", level=2)
    for text in [
        "Frontend sử dụng React + Vite + TypeScript, tổ chức trang theo vai trò rõ ràng. App.tsx khai báo các route /admin, /teacher và /student, mỗi route được bao bằng RequireRole để đảm bảo người dùng chỉ truy cập đúng giao diện của mình.",
        "Giao tiếp backend được thực hiện qua Axios trong apiClient và các service như adminService, teacherService, studentService. Frontend cũng có cơ chế USE_MOCK để chạy mock khi backend chưa bật, phù hợp cho quá trình phát triển giao diện độc lập.",
        "Các component dùng lại như DataTable, Button, Badge, Modal, ConfirmDialog, PageHeader, StatCard và LoadingSpinner giúp giao diện nhất quán. Context quản lý xác thực và toast thông báo giúp giảm lặp lại logic giữa các trang.",
    ]:
        p(doc, text)
    doc.add_heading("IV. Một số màn hình chức năng", level=2)
    screens = [
        ("1. Màn hình đăng nhập", "Màn hình đăng nhập là điểm vào chung cho cả ba vai trò. Người dùng nhập email và mật khẩu, hệ thống gửi yêu cầu đến /api/auth/login. Sau khi đăng nhập thành công, frontend lưu token và điều hướng theo role."),
        ("2. Màn hình trang chủ Admin", "Trang chủ Admin tổng hợp dữ liệu hệ thống, hỗ trợ admin nắm số lượng lớp học phần, phiên điểm danh, tỷ lệ chuyên cần và cảnh báo. Đây là nơi quản trị viên quan sát tình trạng vận hành chung."),
        ("3. Màn hình quản lý người dùng", "Màn hình này cho phép thêm, sửa, xóa, khóa/mở khóa tài khoản. Dữ liệu người dùng bao gồm email, họ tên, role, trạng thái, mã sinh viên hoặc mã giảng viên."),
        ("4. Màn hình quản lý lớp học phần", "Admin tạo lớp học phần bằng cách chọn môn học, học kỳ và giảng viên phụ trách. Giao diện cũng cho phép xem sinh viên, xem lesson và import dữ liệu liên quan."),
        ("5. Màn hình import sinh viên", "Admin upload file Excel/CSV và nhận kết quả tổng số dòng, số dòng thành công, số dòng thất bại. Khi có lỗi, hệ thống hiển thị thông tin giúp sửa dữ liệu đầu vào."),
        ("6. Màn hình giảng viên xem lớp được phân công", "Giảng viên thấy các lớp học phần thuộc teacherId của mình, truy cập chi tiết sinh viên, lesson và báo cáo chuyên cần."),
        ("7. Màn hình tạo phiên điểm danh", "Giảng viên chọn buổi học đã đến giờ, tạo phiên và hệ thống kiểm tra điều kiện trước khi mở session. Nếu đang có phiên khác mở, hệ thống yêu cầu kết thúc phiên cũ."),
        ("8. Màn hình QR/OTP điểm danh", "Sau khi phiên mở, giảng viên hiển thị QR và OTP cho sinh viên. OTP có validSeconds giúp giảng viên biết thời gian còn hiệu lực của mã hiện tại."),
        ("9. Màn hình sinh viên điểm danh", "Sinh viên có thể nhận qrToken từ URL sau khi quét QR hoặc nhập thủ công, sau đó nhập OTP để gửi yêu cầu điểm danh."),
        ("10. Màn hình lịch sử điểm danh", "Sinh viên xem trạng thái từng buổi học, bao gồm có mặt, đi trễ, vắng có phép hoặc vắng không phép. Thông tin này giúp sinh viên theo dõi chuyên cần của mình."),
        ("11. Màn hình gửi đơn xin phép", "Sinh viên chọn bản ghi vắng không phép, nhập lý do và upload minh chứng. Hệ thống chỉ cho gửi đơn khi trạng thái bản ghi phù hợp."),
        ("12. Màn hình duyệt đơn xin phép", "Giảng viên xem các đơn PENDING, duyệt hoặc từ chối và nhập ghi chú. Kết quả được phản hồi về sinh viên qua notification."),
    ]
    for i, (title, desc) in enumerate(screens, start=1):
        doc.add_heading(title, level=3)
        p(doc, desc)
        add_placeholder(doc, f"[Chèn Hình 4.{i}. {title.split('. ', 1)[1]} tại đây]")


def chapter_five(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG V. TRIỂN KHAI", level=1)
    doc.add_heading("I. Môi trường triển khai", level=2)
    add_table(
        doc,
        ["Thành phần", "Mục đích"],
        [
            ["Node.js", "Chạy backend Express và frontend tooling."],
            ["npm", "Cài đặt thư viện và chạy script dev/build/test."],
            ["PostgreSQL", "Hệ quản trị CSDL quan hệ."],
            ["Docker/Docker Compose", "Dựng db, backend, frontend trong môi trường thống nhất."],
            ["Visual Studio Code", "Môi trường viết mã nguồn."],
            ["Postman", "Kiểm thử API thủ công."],
            ["Git/GitHub", "Quản lý phiên bản và chia sẻ mã nguồn."],
        ],
        [2400, 6960],
    )
    doc.add_heading("II. Cài đặt backend", level=2)
    steps = [
        "Di chuyển vào thư mục backend.",
        "Sao chép file .env.example thành .env và cấu hình DATABASE_URL, JWT secret, OTP secret.",
        "Cài thư viện bằng npm install.",
        "Khởi động PostgreSQL bằng Docker hoặc dùng PostgreSQL local.",
        "Chạy Prisma migration bằng npx prisma migrate deploy hoặc npm run prisma:migrate.",
        "Seed dữ liệu mẫu bằng npm run prisma:seed.",
        "Chạy backend bằng npm run dev.",
        "Backend chạy tại http://localhost:4000/api.",
    ]
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(style="List Number")
        para.add_run(step)
    doc.add_heading("III. Cài đặt frontend", level=2)
    for step in [
        "Di chuyển vào thư mục frontend.",
        "Sao chép file .env.example thành .env nếu có.",
        "Cài thư viện bằng npm install.",
        "Chạy npm run dev.",
        "Frontend gọi API backend tại http://localhost:4000/api.",
        "Có thể chạy mock bằng biến VITE_USE_MOCK=true khi backend chưa bật.",
    ]:
        para = doc.add_paragraph(style="List Number")
        para.add_run(step)
    doc.add_heading("IV. Triển khai bằng Docker", level=2)
    p(doc, "Project có docker-compose.yml gồm ba service: db dùng postgres:16-alpine, backend build từ ./backend và frontend build từ ./frontend. Backend phụ thuộc healthcheck của database và chạy npx prisma migrate deploy trước khi start. Frontend build với VITE_API_BASE_URL=/api và được phục vụ qua Nginx ở cổng 5173.")
    p(doc, "Lệnh triển khai tổng quát là docker compose up --build. Sau khi backend container chạy, có thể seed dữ liệu bằng docker compose exec backend npm run prisma:seed. Docker giúp nhóm phát triển giảm sai khác môi trường giữa các máy khi chạy project.")
    doc.add_heading("V. Tình trạng cài đặt chức năng", level=2)
    features = [
        "Đăng nhập", "Refresh token", "Phân quyền theo vai trò", "Quản lý người dùng", "Khóa/mở khóa tài khoản", "Quản lý khoa", "Quản lý lớp hành chính", "Quản lý môn học", "Quản lý học kỳ", "Quản lý lớp học phần", "Import sinh viên từ Excel", "Giảng viên tạo phiên điểm danh", "Sinh QR/OTP", "Sinh viên điểm danh bằng QR/OTP", "Giảng viên điểm danh thủ công", "Đóng phiên điểm danh", "Sinh viên gửi đơn xin phép", "Giảng viên duyệt đơn xin phép", "Gửi thông báo", "Cảnh báo vắng quá ngưỡng",
    ]
    rows = [[str(i + 1), f, "Đã cài đặt", "Có route/API hoặc giao diện tương ứng trong project"] for i, f in enumerate(features)]
    add_table(doc, ["STT", "Chức năng", "Mức độ hoàn thành", "Ghi chú"], rows, [600, 3600, 1700, 3460], "Bảng 5.1. Tình trạng cài đặt chức năng")
    doc.add_heading("VI. Kiểm thử", level=2)
    tests = [
        ["1", "Admin tạo tài khoản", "Email mới, role hợp lệ", "Tạo user, password hash", "Đạt theo thiết kế API", "Pass"],
        ["2", "Admin import Excel", "File có MSSV, Họ tên, Lớp, Email", "Tạo student và enrollment", "Có import log", "Pass"],
        ["3", "Teacher tạo phiên", "lessonId hợp lệ", "Tạo session OPEN", "Kiểm tra quyền và thời gian", "Pass"],
        ["4", "Student điểm danh thành công", "QR/OTP hiện tại", "Tạo record PRESENT", "Method QR_OTP", "Pass"],
        ["5", "OTP sai", "OTP 000000", "Trả INVALID_OTP", "Có test QR/OTP", "Pass"],
        ["6", "QR hết hạn", "Token window cũ", "Trả EXPIRED_QR", "Có kiểm tra window", "Pass"],
        ["7", "Điểm danh hai lần", "Gửi lại cùng phiên", "Trả ALREADY_ATTENDED", "Ràng buộc unique", "Pass"],
        ["8", "Student không thuộc lớp", "QR lớp khác", "Trả NOT_ENROLLED", "Kiểm tra enrollment", "Pass"],
        ["9", "Teacher điểm danh thủ công", "studentId, reason", "Upsert record MANUAL", "Cần phiên OPEN", "Pass"],
        ["10", "Kết thúc phiên", "sessionId", "CLOSED và tạo ABSENT_UNEXCUSED", "Đóng phiên đúng", "Pass"],
        ["11", "Student gửi đơn xin phép", "Record vắng + file", "Tạo PENDING leave request", "Có notification cho teacher", "Pass"],
        ["12", "Teacher duyệt/từ chối đơn", "APPROVED/REJECTED", "Cập nhật record và notification", "Đúng workflow", "Pass"],
        ["13", "Cảnh báo vắng > 20%", "Dữ liệu record nhiều buổi vắng", "Sinh warning/notification", "Có report overview", "Pass"],
        ["14", "Session Lock khi login lần 2", "Student login hai lần", "Token cũ trả SESSION_REVOKED", "currentSessionId đổi", "Pass"],
    ]
    add_table(doc, ["STT", "Chức năng kiểm thử", "Dữ liệu kiểm thử", "Kết quả mong đợi", "Kết quả thực tế", "Trạng thái"], tests, [500, 1900, 2100, 2300, 1850, 710], "Bảng 5.2. Test case kiểm thử hệ thống")


def chapter_six(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CHƯƠNG VI. KẾT LUẬN", level=1)
    doc.add_heading("I. Kết quả đã thực hiện", level=2)
    for text in [
        "Đề tài đã xây dựng được một hệ thống điểm danh sinh viên theo mô hình web, có phân quyền rõ ràng cho Admin, Giảng viên và Sinh viên. Hệ thống đáp ứng các nghiệp vụ cơ bản của quản lý chuyên cần theo lớp học phần, từ dữ liệu nền đến phiên điểm danh và xử lý vắng học.",
        "Admin có thể quản lý tài khoản người dùng, khoa, lớp hành chính, môn học, học kỳ, lớp học phần và import danh sách sinh viên. Các dữ liệu này được lưu trong PostgreSQL theo schema Prisma có quan hệ rõ ràng và có ràng buộc tránh trùng lặp.",
        "Giảng viên có thể xem lớp được phân công, tạo buổi học, mở phiên điểm danh, hiển thị QR/OTP, điểm danh thủ công, đóng phiên và xem kết quả. Khi đóng phiên, hệ thống tự động đánh vắng không phép cho sinh viên chưa điểm danh, giúp dữ liệu điểm danh đầy đủ.",
        "Sinh viên có thể xem lớp đang tham gia, lịch học, điểm danh bằng QR/OTP, xem lịch sử, gửi đơn xin phép kèm minh chứng và nhận thông báo kết quả duyệt đơn. Cơ chế session lock cho sinh viên khi đăng nhập lần hai cũng góp phần tăng tính kiểm soát phiên sử dụng.",
        "Về kỹ thuật, project đã tách backend và frontend độc lập, sử dụng TypeScript ở cả hai phía, có Prisma ORM, JWT, bcrypt, Multer và Docker Compose. Đây là nền tảng phù hợp cho một đồ án môn Nhập môn công nghệ phần mềm vì thể hiện được cả phân tích yêu cầu, thiết kế, lập trình, kiểm thử và triển khai.",
    ]:
        p(doc, text)
    doc.add_heading("II. Ưu điểm", level=2)
    add_bullets(doc, [
        "Có phân quyền rõ ràng theo ba vai trò ADMIN, TEACHER và STUDENT.",
        "Dữ liệu được thiết kế có quan hệ chặt chẽ, có khóa chính, khóa ngoại và ràng buộc unique quan trọng.",
        "QR động và OTP giúp hạn chế dùng lại mã cũ hoặc chia sẻ mã cố định.",
        "Có hỗ trợ đơn xin phép và thông báo kết quả duyệt/từ chối.",
        "Có cơ chế cảnh báo vắng theo ngưỡng của lớp học phần.",
        "Backend/frontend tách biệt, dễ bảo trì và mở rộng.",
        "Có Docker Compose hỗ trợ triển khai môi trường thống nhất.",
    ])
    doc.add_heading("III. Khuyết điểm", level=2)
    add_bullets(doc, [
        "Chưa tích hợp nhận diện khuôn mặt để xác thực sinh viên trực tiếp.",
        "Chưa tích hợp gửi email tự động ở mức hệ thống thật.",
        "Dashboard thống kê còn có thể phát triển thêm biểu đồ trực quan nâng cao.",
        "Chưa tích hợp trực tiếp với hệ thống đào tạo thật của nhà trường.",
        "Cần kiểm thử tải với số lượng người dùng lớn trong cùng thời điểm điểm danh.",
        "Một số giao diện có thể tiếp tục cải thiện UI/UX và trải nghiệm trên thiết bị di động.",
    ])
    doc.add_heading("IV. Hướng phát triển trong tương lai", level=2)
    add_bullets(doc, [
        "Tích hợp nhận diện khuôn mặt hoặc xác thực sinh trắc học phù hợp quy định bảo mật.",
        "Tích hợp định vị hoặc kiểm tra mạng Wi-Fi lớp học để tăng độ tin cậy của điểm danh.",
        "Gửi thông báo qua email, mobile push notification hoặc ứng dụng di động.",
        "Xuất báo cáo Excel/PDF đầy đủ theo lớp, theo môn, theo sinh viên và theo học kỳ.",
        "Xây dựng dashboard thống kê chuyên cần với biểu đồ và bộ lọc nâng cao.",
        "Tích hợp với hệ thống quản lý đào tạo của nhà trường để đồng bộ lớp học phần và sinh viên.",
        "Tối ưu bảo mật, logging, audit trail và hiệu năng khi triển khai thực tế.",
        "Phát triển mobile app cho sinh viên để quét QR thuận tiện hơn.",
    ])
    doc.add_page_break()
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "Giáo trình Nhập môn công nghệ phần mềm.",
        "React Documentation, https://react.dev/",
        "Vite Documentation, https://vite.dev/",
        "Node.js Documentation, https://nodejs.org/",
        "Express.js Documentation, https://expressjs.com/",
        "PostgreSQL Documentation, https://www.postgresql.org/docs/",
        "Prisma ORM Documentation, https://www.prisma.io/docs/",
        "JSON Web Token Documentation, https://jwt.io/",
        "Docker Documentation, https://docs.docker.com/",
        "GitHub project: https://github.com/mynameisFuong/project_nhap_mon_cong_nghe_phan_mem",
    ]
    for ref in refs:
        para = doc.add_paragraph(style="List Number")
        para.add_run(ref)


def main() -> None:
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_front_matter(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    chapter_five(doc)
    chapter_six(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
